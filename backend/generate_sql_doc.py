from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_sql_document():
    doc = Document()

    title = doc.add_heading('Consultas SQL del Backend', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Sistema de Gesti\u00f3n EPS - Base de Datos 2026-01')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # ==================== SECCION 1: REPORTES ====================
    doc.add_heading('1. Consultas de Reportes y Dashboard', level=1)

    reports = [
        {
            'title': '1.1 Afiliados con Facturas Pendientes',
            'file': 'backend/app.py (l\u00ednea 1157-1166)',
            'query': """SELECT a.id_afiliado, u.nombre, u.email,
       COUNT(f.id_factura) as facturas_pendientes,
       SUM(f.total) as total_pendiente
FROM afiliados a
JOIN usuarios u ON a.id_usuario = u.id_usuario
LEFT JOIN facturas f ON a.id_afiliado = f.id_afiliado AND f.estado = 'PENDIENTE'
GROUP BY a.id_afiliado, u.nombre, u.email
HAVING COUNT(f.id_factura) > 0;""",
            'description': 'Lista todos los afiliados que tienen facturas pendientes, mostrando nombre, email, cantidad de facturas y total pendiente.'
        },
        {
            'title': '1.2 Facturaci\u00f3n por Plan',
            'file': 'backend/app.py (l\u00ednea 1185-1193)',
            'query': """SELECT p.nombre, COUNT(DISTINCT a.id_afiliado) as total_afiliados,
       COALESCE(SUM(f.total), 0) as facturacion_total
FROM planes p
LEFT JOIN afiliados a ON p.id_plan = a.id_plan
LEFT JOIN facturas f ON a.id_afiliado = f.id_afiliado AND f.estado = 'PAGADA'
GROUP BY p.id_plan, p.nombre;""",
            'description': 'Muestra el total de facturaci\u00f3n agrupado por plan de salud, incluyendo cantidad de afiliados y total de facturas pagadas.'
        },
        {
            'title': '1.3 Diagn\u00f3sticos M\u00e1s Frecuentes',
            'file': 'backend/app.py (l\u00ednea 1210-1219)',
            'query': """SELECT p.especialidad, h.diagnostico, COUNT(*) as cantidad
FROM historial_clinico h
JOIN profesionales p ON h.id_profesional = p.id_profesional
GROUP BY p.especialidad, h.diagnostico
ORDER BY cantidad DESC
LIMIT 20;""",
            'description': 'Lista los 20 diagn\u00f3sticos m\u00e1s comunes, organizados por especialidad m\u00e9dica.'
        },
        {
            'title': '1.4 Centros de Salud M\u00e1s Utilizados',
            'file': 'backend/app.py (l\u00ednea 1236-1244)',
            'query': """SELECT c.nombre, c.ciudad, COUNT(cit.id_cita) as total_citas
FROM centros_salud c
LEFT JOIN citas cit ON c.id_centro = cit.id_centro
GROUP BY c.id_centro, c.nombre, c.ciudad
ORDER BY total_citas DESC;""",
            'description': 'Ranking de centros de salud por cantidad de citas m\u00e9dicas atendidas, ordenado de mayor a menor.'
        },
        {
            'title': '1.5 Total Facturas Pagadas',
            'file': 'backend/app.py (l\u00ednea 1299-1301)',
            'query': 'SELECT COALESCE(SUM(total), 0) FROM facturas WHERE estado = "PAGADA"',
            'description': 'Calcula la suma total de todas las facturas con estado PAGADA.'
        },
        {
            'title': '1.6 Total Facturas Pendientes',
            'file': 'backend/app.py (l\u00ednea 1303-1305)',
            'query': 'SELECT COALESCE(SUM(total), 0) FROM facturas WHERE estado = "PENDIENTE"',
            'description': 'Calcula la suma total de todas las facturas con estado PENDIENTE.'
        },
        {
            'title': '1.7 Afiliados por Plan',
            'file': 'backend/app.py (l\u00ednea 1307-1314)',
            'query': """SELECT p.nombre, COUNT(a.id_afiliado) as cantidad
FROM planes p
LEFT JOIN afiliados a ON p.id_plan = a.id_plan
GROUP BY p.id_plan, p.nombre;""",
            'description': 'Cuenta la cantidad de afiliados asociados a cada plan de salud.'
        }
    ]

    for report in reports:
        doc.add_heading(report['title'], level=2)
        doc.add_paragraph(f'Archivo: {report["file"]}').runs[0].italic = True

        p = doc.add_paragraph()
        p.add_run('Descripci\u00f3n: ').bold = True
        p.add_run(report['description'])

        doc.add_paragraph('')
        p = doc.add_paragraph()
        runner = p.add_run('Consulta SQL:')
        runner.bold = True

        code_para = doc.add_paragraph(report['query'])
        code_para.style = 'No Spacing'
        code_para.paragraph_format.left_indent = Inches(0.3)

        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 100, 0)

        doc.add_paragraph('')

    # ==================== SECCION 2: SCRIPTS SQL ====================
    doc.add_heading('2. Scripts SQL de Configuraci\u00f3n', level=1)

    # Seed SQL
    doc.add_heading('2.1 seed.sql - Datos Iniciales', level=2)

    seed_queries = [
        ('Insertar Usuarios Default', 'backend/scripts/seed.sql (l\u00edneas 7-11)', """INSERT IGNORE INTO usuarios (email, password, nombre, rol) VALUES
('admin@eps.com', '\$2b\$12\$...', 'Administrador', 'ADMIN'),
('afiliado@test.com', '\$2b\$12\$...', 'Juan P\u00e9rez', 'AFILIADO'),
('profesional@test.com', '\$2b\$12\$...', 'Dra. Mar\u00eda Garc\u00eda', 'PROFESIONAL'),
('drlorenzo@test.com', '\$2b\$12\$...', 'Dr. Carlos Lorenzo', 'PROFESIONAL');"""),
        ('Crear Afiliado de Prueba', 'backend/scripts/seed.sql (l\u00edneas 13-15)', """INSERT IGNORE INTO afiliados (id_usuario, documento, telefono, direccion, id_plan)
SELECT id_usuario, '12345678', '3001234567', 'Calle 123, Ciudad', 1
FROM usuarios WHERE email = 'afiliado@test.com';"""),
        ('Crear Profesionales de Prueba', 'backend/scripts/seed.sql (l\u00edneas 17-23)', """INSERT IGNORE INTO profesionales (id_usuario, especialidad, id_centro)
SELECT id_usuario, 'Medicina General', 1 FROM usuarios WHERE email = 'profesional@test.com';
INSERT IGNORE INTO profesionales (id_usuario, especialidad, id_centro)
SELECT id_usuario, 'Cardiolog\u00eda', 1 FROM usuarios WHERE email = 'drlorenzo@test.com';"""),
        ('Insertar Disponibilidad Profesional', 'backend/scripts/seed.sql (l\u00edneas 25-39)', """INSERT IGNORE INTO disponibilidad_profesional (id_profesional, dia_semana, hora_inicio, hora_fin, activo)
SELECT p.id_profesional, dia_semana, hora_inicio, hora_fin, 1
FROM profesionales p
CROSS JOIN (
    SELECT 1 AS dia_semana, '08:00' AS hora_inicio, '17:00' AS hora_fin
    UNION ALL SELECT 2, '08:00', '17:00'
    UNION ALL SELECT 3, '08:00', '17:00'
    UNION ALL SELECT 4, '08:00', '17:00'
    UNION ALL SELECT 5, '08:00', '17:00'
) AS horarios
WHERE p.usuario_id IN (SELECT id_usuario FROM usuarios WHERE email IN (...))
AND NOT EXISTS (
    SELECT 1 FROM disponibilidad_profesional dp
    WHERE dp.id_profesional = p.id_profesional AND dp.dia_semana = horarios.dia_semana
);""")
    ]

    for name, file, query in seed_queries:
        doc.add_heading(name, level=3)
        doc.add_paragraph(f'Archivo: {file}').runs[0].italic = True
        code_para = doc.add_paragraph(query)
        code_para.style = 'No Spacing'
        code_para.paragraph_format.left_indent = Inches(0.3)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        doc.add_paragraph('')

    # Limpiar duplicados
    doc.add_heading('2.2 limpiar_duplicados.sql - Eliminaci\u00f3n de Duplicados', level=2)

    dup_queries = [
        ('Eliminar Duplicados por Documento', 'backend/scripts/limpiar_duplicados.sql (l\u00edneas 5-8)', """DELETE a1 FROM afiliados a1
JOIN afiliados a2
WHERE a1.id_afiliado > a2.id_afiliado
AND a1.documento = a2.documento;"""),
        ('Eliminar Duplicados por Tel\u00e9fono', 'backend/scripts/limpiar_duplicados.sql (l\u00edneas 10-15)', """DELETE a1 FROM afiliados a1
JOIN afiliados a2
WHERE a1.id_afiliado > a2.id_afiliado
AND a1.telefono = a2.telefono
AND a1.telefono IS NOT NULL AND a2.telefono IS NOT NULL;"""),
        ('Agregar Constraint \u00danico Documento', 'backend/scripts/limpiar_duplicados.sql (l\u00ednea 19)', 'ALTER TABLE afiliados ADD CONSTRAINT IF NOT EXISTS unique_documento UNIQUE (documento);'),
        ('Agregar Constraint \u00danico Tel\u00e9fono', 'backend/scripts/limpiar_duplicados.sql (l\u00ednea 22)', 'ALTER TABLE afiliados ADD CONSTRAINT IF NOT EXISTS unique_telefono UNIQUE (telefono);')
    ]

    for name, file, query in dup_queries:
        doc.add_heading(name, level=3)
        doc.add_paragraph(f'Archivo: {file}').runs[0].italic = True
        code_para = doc.add_paragraph(query)
        code_para.style = 'No Spacing'
        code_para.paragraph_format.left_indent = Inches(0.3)
        for run in code_para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        doc.add_paragraph('')

    # ==================== SECCION 3: ESQUEMA ====================
    doc.add_heading('3. Esquema de Base de Datos (database.sql)', level=1)

    tables = [
        ('usuarios', 'Usuarios del sistema', 'ADMIN, AFILIADO, PROFESIONAL'),
        ('planes', 'Planes de salud disponibles', ''),
        ('afiliados', 'Miembros afiliados a la EPS', ''),
        ('centros_salud', 'Centros de atenci\u00f3n m\u00e9dica', ''),
        ('profesionales', 'Profesionales de la salud', ''),
        ('facturas', 'Facturas de los afiliados', ''),
        ('pagos', 'Pagos de facturas', ''),
        ('citas', 'Citas m\u00e9dicas', ''),
        ('historial_clinico', 'Registros cl\u00ednicos de afiliados', ''),
        ('disponibilidad_profesional', 'Horarios de disponibilidad', '')
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tabla'
    hdr_cells[1].text = 'Descripci\u00f3n'
    hdr_cells[2].text = 'Notas'

    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '2E86AB')

    for name, desc, notes in tables:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = desc
        row_cells[2].text = notes

    doc.add_paragraph('')

    # ==================== SECCION 4: ORM QUERIES ====================
    doc.add_heading('4. Patrones de Consultas ORM', level=1)

    orm_queries = [
        ('B\u00fasqueda de usuario por email', 'Usuario.query.filter_by(email=...)'),
        ('B\u00fasqueda de afiliado por id_usuario', 'Afiliado.query.filter_by(id_usuario=...)'),
        ('B\u00fasqueda de afiliado por documento', 'Afiliado.query.filter_by(documento=..., tipo_documento=...)'),
        ('B\u00fasqueda de afiliado por tel\u00e9fono', 'Afiliado.query.filter_by(telefono=...)'),
        ('B\u00fasqueda de profesional por id_usuario', 'Profesional.query.filter_by(id_usuario=...)'),
        ('Obtener disponibilidad profesional', 'DisponibilidadProfesional.query.filter_by(...)'),
        ('CRUD de citas m\u00e9dicas', 'Cita.query.filter_by(...) / .add() / .update() / .delete()'),
        ('CRUD de facturas', 'Factura.query.filter_by(...) / .add() / .update() / .delete()'),
        ('CRUD de pagos', 'Pago.query.filter_by(...) / .add() / .delete()'),
        ('CRUD de historial cl\u00ednico', 'HistorialClinico.query.filter_by(...) / .add() / .delete()')
    ]

    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'

    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Operaci\u00f3n'
    hdr2[1].text = 'Patr\u00f3n ORM'

    for cell in hdr2:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '2E86AB')

    for op, pattern in orm_queries:
        row = table2.add_row().cells
        row[0].text = op
        row[1].text = pattern

    doc.add_paragraph('')

    # ==================== SECCION 5: CONFIG ====================
    doc.add_heading('5. Configuraci\u00f3n de Conexi\u00f3n', level=1)

    doc.add_paragraph('Archivo: backend/config.py')
    config_query = doc.add_paragraph('SQLALCHEMY_DATABASE_URI = f\'mysql+pymysql://{DB_USER}:{DB_PASSWORD}@{DB_HOST}:{DB_PORT}/{DB_NAME}?charset=utf8mb4\'')
    config_query.style = 'No Spacing'
    config_query.paragraph_format.left_indent = Inches(0.3)
    for run in config_query.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph('')

    # Footer
    doc.add_paragraph('')
    footer = doc.add_paragraph('Documento generado autom\u00e1ticamente - Proyecto Final Base de Datos 2026-01')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True

    output_path = r'C:\Users\Asus\OneDrive\Desktop\Universidad\Base de Datos 2026-01\Proyecto Final\Diagramas\Consultas_SQL_Backend.docx'
    doc.save(output_path)
    print(f'Documento guardado en: {output_path}')

if __name__ == '__main__':
    create_sql_document()